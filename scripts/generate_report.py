import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_report():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- COVER PAGE ---
    # (Simplified for the generator)
    title = doc.add_heading('THREE COMPOSITE BARS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Computational Mechanics & Structural Analysis Platform')
    run.font.size = Pt(14)
    run.italic = True

    doc.add_paragraph('\n' * 5)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
    run.font.size = Pt(12)

    doc.add_paragraph('\n' * 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Developer: SENAPATHI YASWANTH (RA17)')
    run.bold = True
    run.font.size = Pt(12)

    doc.add_page_break()

    # --- 1. OVERVIEW ---
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        "The 'Three Composite Bars' platform is a specialized engineering tool designed to solve "
        "statically indeterminate structural systems involving parallel bars of different materials. "
        "The project integrates advanced numerical methods with a modern full-stack web architecture."
    )

    # --- 2. PREMIUM FEATURES [NEW] ---
    doc.add_heading('2. Advanced Engineering Features', level=1)
    
    features = [
        ('Structural Mass Estimation', 'Calculates individual rod mass and total system weight based on material density ($kg/m^3$).'),
        ('Safety Analysis Dashboard', 'Identifies the critical component carrying the highest stress intensity.'),
        ('Step-by-Step Derivation', 'Provides a complete mathematical breakdown of the solving process using Compatibility and Equilibrium conditions.'),
        ('AI Problem Parsing', 'Allows natural language input for structural parameters, powered by hybrid NLP logic.'),
        ('Dynamic Sensitivity Analysis', 'Visualizes stress redistribution across the system as component geometries vary.')
    ]
    
    for f_title, f_desc in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{f_title}: ')
        run.bold = True
        p.add_run(f_desc)

    # --- 3. MATHEMATICAL ENGINE ---
    doc.add_heading('3. Mathematical Foundation', level=1)
    doc.add_paragraph(
        "The solver engine is built on two core mechanical principles:"
    )
    
    p = doc.add_paragraph()
    p.add_run('1. Equilibrium: ').bold = True
    p.add_run('The sum of internal forces equals the total applied load ($P_1 + P_2 + ... = P_{total}$).')
    
    p = doc.add_paragraph()
    p.add_run('2. Compatibility: ').bold = True
    p.add_run('All parallel bars connected to rigid plates must undergo the same deformation ($\\delta_{common}$).')

    # --- 4. SYSTEM ARCHITECTURE ---
    doc.add_heading('4. Technical Stack', level=1)
    tech = doc.add_table(rows=1, cols=2)
    tech.style = 'Table Grid'
    hdr_cells = tech.rows[0].cells
    hdr_cells[0].text = 'Layer'
    hdr_cells[1].text = 'Technology'
    
    stack = [
        ('Frontend', 'React 18, Vite, Framer Motion, Recharts'),
        ('Backend', 'FastAPI, Uvicorn, Python 3.12+'),
        ('Numerical', 'NumPy (Matrix Algebra)'),
        ('Database', 'MongoDB, Motor (Async Driver)')
    ]
    
    for layer, tech_name in stack:
        row_cells = tech.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = tech_name

    doc.add_paragraph('\n')
    doc.add_paragraph(
        "The project is hardened with professional error handling (503 Service Unavailable) "
        "and robust JSON serialization for high-precision numerical outputs."
    )

    # --- SAVE ---
    output_path = os.path.join('reports', 'PROJECT_DOCUMENTATION.docx')
    os.makedirs('reports', exist_ok=True)
    doc.save(output_path)
    print(f"Report generated successfully at {output_path}")

if __name__ == "__main__":
    create_report()
